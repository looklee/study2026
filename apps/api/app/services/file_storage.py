# 文件上传与存储服务

import os
import uuid
import shutil
from pathlib import Path
from typing import Optional, Dict, Any, List
from fastapi import UploadFile, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
import aiofiles
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models import KnowledgeDocument
from app.core.config import settings


# 存储配置
UPLOAD_DIR = Path(settings.BASE_DIR) / "uploads"
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
ALLOWED_EXTENSIONS = {
    # 文档
    "pdf", "doc", "docx", "txt", "md", "rtf",
    # 图片
    "jpg", "jpeg", "png", "gif", "webp", "svg",
    # 代码
    "py", "js", "ts", "java", "cpp", "c", "h", "html", "css", "json", "yaml", "yml",
    # 数据
    "csv", "xlsx", "xls",
    # 其他
    "zip", "rar", "7z"
}

MIME_TYPES = {
    "pdf": "application/pdf",
    "doc": "application/msword",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
    "md": "text/markdown",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "gif": "image/gif",
    "webp": "image/webp",
    "svg": "image/svg+xml",
    "py": "text/x-python",
    "js": "application/javascript",
    "ts": "application/typescript",
    "json": "application/json",
    "csv": "text/csv",
    "zip": "application/zip",
}


def get_file_category(file_extension: str) -> str:
    """获取文件分类"""
    ext = file_extension.lower()
    if ext in ["pdf", "doc", "docx", "txt", "md", "rtf"]:
        return "document"
    elif ext in ["jpg", "jpeg", "png", "gif", "webp", "svg"]:
        return "image"
    elif ext in ["py", "js", "ts", "java", "cpp", "c", "h", "html", "css"]:
        return "code"
    elif ext in ["csv", "xlsx", "xls"]:
        return "data"
    else:
        return "other"


def validate_file(file: UploadFile) -> Dict[str, Any]:
    """验证文件"""
    # 获取文件扩展名
    filename = file.filename or ""
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    
    if not ext:
        raise HTTPException(status_code=400, detail="无法识别文件类型")
    
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型：.{ext}。允许的类型：{', '.join(sorted(ALLOWED_EXTENSIONS)[:10])}..."
        )
    
    return {
        "filename": filename,
        "extension": ext,
        "category": get_file_category(ext),
        "mime_type": MIME_TYPES.get(ext, "application/octet-stream")
    }


async def save_file(
    file: UploadFile,
    user_id: int,
    category: str = "general",
    db: AsyncSession = None
) -> Dict[str, Any]:
    """保存上传的文件"""
    import aiofiles
    
    # 验证文件
    file_info = validate_file(file)
    
    # 生成唯一文件名
    file_ext = file_info["extension"]
    unique_filename = f"{uuid.uuid4().hex}.{file_ext}"
    
    # 创建目录结构：uploads/{user_id}/{category}/
    user_dir = UPLOAD_DIR / str(user_id) / category
    user_dir.mkdir(parents=True, exist_ok=True)
    
    file_path = user_dir / unique_filename
    
    # 保存文件
    file_size = 0
    async with aiofiles.open(file_path, "wb") as out_file:
        while True:
            chunk = await file.read(1024 * 1024)  # 每次读取 1MB
            if not chunk:
                break
            file_size += len(chunk)
            
            # 检查文件大小
            if file_size > MAX_FILE_SIZE:
                file_path.unlink(missing_ok=True)
                raise HTTPException(
                    status_code=400,
                    detail=f"文件过大 (最大 {MAX_FILE_SIZE // 1024 // 1024}MB)"
                )
            
            await out_file.write(chunk)
    
    # 保存到数据库
    doc = None
    if db:
        doc = KnowledgeDocument(
            user_id=user_id,
            file_name=file_info["filename"],
            file_size=file_size,
            mime_type=file_info["mime_type"],
            category=category,
            status="uploaded"
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)
    
    return {
        "id": doc.id if doc else None,
        "file_name": file_info["filename"],
        "file_path": str(file_path),
        "file_size": file_size,
        "mime_type": file_info["mime_type"],
        "category": category,
        "url": f"/api/files/{user_id}/{category}/{unique_filename}"
    }


async def get_file(file_id: int, db: AsyncSession) -> Optional[KnowledgeDocument]:
    """获取文件信息"""
    result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == file_id)
    )
    return result.scalar_one_or_none()


async def delete_file(file_id: int, user_id: int, db: AsyncSession) -> bool:
    """删除文件"""
    doc = await get_file(file_id, db)
    
    if not doc or doc.user_id != user_id:
        return False
    
    # 删除物理文件
    file_path = Path(doc.file_path) if doc.file_path else None
    if file_path and file_path.exists():
        file_path.unlink()
    
    # 从数据库删除
    await db.delete(doc)
    await db.commit()
    
    return True


async def list_user_files(
    user_id: int,
    category: str = None,
    db: AsyncSession = None
) -> List[Dict[str, Any]]:
    """列出用户的所有文件"""
    from sqlalchemy import select
    
    query = select(KnowledgeDocument).where(KnowledgeDocument.user_id == user_id)
    
    if category:
        query = query.where(KnowledgeDocument.category == category)
    
    result = await db.execute(query.order_by(KnowledgeDocument.created_at.desc()))
    docs = result.scalars().all()
    
    return [
        {
            "id": doc.id,
            "file_name": doc.file_name,
            "file_size": doc.file_size,
            "mime_type": doc.mime_type,
            "category": doc.category,
            "status": doc.status,
            "created_at": doc.created_at.isoformat() if doc.created_at else None
        }
        for doc in docs
    ]


async def extract_text_from_file(file_path: str) -> str:
    """从文件提取文本内容"""
    path = Path(file_path)
    ext = path.suffix.lower()
    
    try:
        # 文本文件直接读取
        if ext in [".txt", ".md", ".py", ".js", ".ts", ".json", ".yaml", ".yml", ".csv"]:
            async with aiofiles.open(path, "r", encoding="utf-8") as f:
                return await f.read()
        
        # PDF 文件 (需要 pdfplumber 或 PyPDF2)
        elif ext == ".pdf":
            try:
                import PyPDF2
                text = ""
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text()
                return text
            except ImportError:
                return "[PDF 内容需要安装 PyPDF2]"
        
        # Word 文档
        elif ext in [".docx"]:
            try:
                from docx import Document
                doc = Document(path)
                return "\n".join([p.text for p in doc.paragraphs])
            except ImportError:
                return "[DOCX 内容需要安装 python-docx]"
        
        return ""
    except Exception as e:
        return f"[文件读取失败：{str(e)}]"
